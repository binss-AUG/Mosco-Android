import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

DOCX_TEMPLATE = r'D:\MEox\UITer\DOAN\Mosco_Megre\Mosco\uit_phu_luc_3_mau_bao_cao.docx'
DOCX_OUTPUT = r'D:\MEox\UITer\DOAN\Mosco_Megre\Mosco\uit_phu_luc_3_mau_bao_cao_filled.docx'
CH3_PATH = r'D:\MEox\UITer\DOAN\Mosco_Megre\Mosco\raw-doc\03_chuong_3_phan_tich_thiet_ke.md'
CH4_PATH = r'D:\MEox\UITer\DOAN\Mosco_Megre\Mosco\raw-doc\04_chuong_4_trien_khai_hieu_nang.md'

def load_markdown_lines(path):
    with open(path, 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()
    if lines and lines[0].startswith('\ufeff'):
        lines[0] = lines[0][1:]
    return lines

def is_heading(line):
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    return m

def is_table_row(line):
    return line.strip().startswith('|')

def is_code_block(line):
    return line.strip().startswith('```')

def is_horizontal_rule(line):
    s = line.strip().lstrip('\\')
    return s.startswith('---')

def style_for_heading(level):
    if level == 1: return 'Heading 1'
    if level == 2: return 'Heading 2'
    if level == 3: return 'Heading 3'
    if level == 4: return 'Heading 4'
    return 'Normal'

def heading_level_from_markdown(hash_count):
    """Map markdown heading level to Word heading level.
       #  -> Heading 1  (chapter title)
       ## -> Heading 2  (major section)
       ### -> Heading 3 (sub-section)
       #### -> Heading 4 (sub-sub-section)
    """
    return hash_count

def process_content(doc, lines):
    heading_start_re = re.compile(r'^#{2,4}\s+')  # only h2-h4
    code_block = False
    table_rows = []
    
    for raw_line in lines:
        stripped = raw_line.rstrip('\n').rstrip('\r')
        
        if is_code_block(raw_line):
            code_block = not code_block
            continue
        if code_block:
            continue
        
        if not stripped:
            continue
        
        if is_horizontal_rule(raw_line):
            continue
        
        m = is_heading(stripped)
        if m:
            hashes = len(m.group(1))
            text = m.group(2)
            level = heading_level_from_markdown(hashes)
            p = doc.add_paragraph(text, style=f'Heading {level}')
            continue
        
        if is_table_row(raw_line):
            table_rows.append(raw_line)
            continue
        else:
            if table_rows:
                _add_table(doc, table_rows)
                table_rows = []
        
        p = doc.add_paragraph(stripped, style='Normal')
    
    if table_rows:
        _add_table(doc, table_rows)

def _add_table(doc, rows):
    cells_data = []
    for row in rows:
        parts = [c.strip() for c in row.strip().split('|')[1:-1]]
        if all(p.replace('-','').strip()=='' for p in parts):
            continue
        cells_data.append(parts)
    
    if not cells_data:
        return
    
    header = cells_data[0]
    data = cells_data[1:]
    cols = max(len(r) for r in cells_data) if cells_data else 0
    if cols == 0:
        return
    
    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.style = 'Light Grid Accent 1'
    
    for ci, h in enumerate(header):
        if ci < cols:
            cell = table.rows[0].cells[ci]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            if ci < cols:
                table.rows[ri+1].cells[ci].text = val
    
    doc.add_paragraph()  # spacing after table

# -------------------------------------------------------------
def main():
    doc = Document(DOCX_TEMPLATE)
    body = doc.element.body
    paras = body.findall(qn('w:p'))
    
    # Find boundaries
    ch3_start_idx = None
    ref_idx = None
    # Also find the heading "Chương 4" or the placeholder before TÀI LIỆU THAM KHẢO
    ch4_placeholder_end = None
    
    for i, p in enumerate(paras):
        text = p.text.strip() if p.text else ''
        if 'Dựa trên cơ sở lý thuyết' in text:
            ch3_start_idx = i
        if 'TÀI LIỆU THAM KHẢO' in text:
            ref_idx = i
        if 'Chủ đề cấp độ 2' in text and ch3_start_idx is not None and ref_idx is None:
            ch4_placeholder_end = i  # last placeholder before references
    
    if ch3_start_idx is None or ref_idx is None:
        print("ERROR: Could not find boundaries")
        print(f"ch3_start_idx={ch3_start_idx}, ref_idx={ref_idx}")
        return
    
    # Remove placeholder content (from ch3_start_idx to ref_idx-1 inclusive)
    paras_to_remove = list(paras[ch3_start_idx:ref_idx])
    for p in paras_to_remove:
        p.getparent().remove(p)
    
    # Now insert new content before the first remaining element after the removed range
    # We need to find the insertion point - the first element after the removed range
    # Since we removed elements, the body structure changed. We need to find 
    # the reference element (TÀI LIỆU THAM KHẢO) again after removal.
    
    # Find reference paragraph in the current body
    current_paras = body.findall(qn('w:p'))
    insert_before_para = None
    for p in current_paras:
        text = p.text.strip() if p.text else ''
        if 'TÀI LIỆU THAM KHẢO' in text:
            insert_before_para = p
            break
    
    if insert_before_para is None:
        print("ERROR: Could not find TÀI LIỆU THAM KHẢO after removal")
        return
    
    # Build content from markdown
    ch3_lines = load_markdown_lines(CH3_PATH)
    ch4_lines = load_markdown_lines(CH4_PATH)
    
    # Add content temporarily to doc end, then move them
    # Actually, better approach: add content to a temp document, 
    # extract their XML, and insert before the reference.
    
    from docx import Document as Doc
    temp_doc = Doc()
    
    # Process Ch3 content (markdown includes # chapter title)
    process_content(temp_doc, ch3_lines)
    
    # Process Ch4 content (markdown includes # chapter title)
    process_content(temp_doc, ch4_lines)
    
    # Move all children of temp_body (paragraphs, tables, etc.) in order
    temp_body = temp_doc.element.body
    children = list(temp_body)
    for child in children:
        insert_before_para.addprevious(child)
    
    doc.save(DOCX_OUTPUT)
    print("SUCCESS: Document saved!")

if __name__ == '__main__':
    main()
